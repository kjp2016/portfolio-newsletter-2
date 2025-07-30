#app.py
import streamlit as st
import pandas as pd
import json
import os
from datetime import datetime, timedelta
import logging
from typing import Dict, List, Tuple, Optional, Any
import re
from io import BytesIO

# Document processing imports
import PyPDF2
import docx
import openpyxl
from openai import OpenAI
import yfinance as yf

# --- Corrected Imports ---
# Imports the necessary functions from your other project files
from portfolio_analysis import get_batch_stock_data
from google_sheets_storage import (
    init_google_sheet,
    save_user_portfolio_to_sheets,
    get_user_portfolio_from_sheets,
    get_all_users_from_sheets
)
# This now correctly imports the newsletter function from main.py
from main import generate_newsletter_for_user


# ---------- Configuration ----------
OPENAI_MODEL = "gpt-4o-mini"
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# ---------- Logging ----------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-7s %(message)s",
    datefmt="%H:%M:%S"
)

# ---------- Document Processing Functions ----------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from Word document."""
    try:
        doc = docx.Document(BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text
    except Exception as e:
        logging.error(f"Error extracting text from DOCX: {e}")
        return ""

def extract_data_from_excel(file_bytes: bytes) -> pd.DataFrame:
    """Extract data from Excel file."""
    try:
        logging.info("Starting Excel file extraction...")
        excel_file = pd.ExcelFile(BytesIO(file_bytes))
        logging.info(f"Excel file sheets: {excel_file.sheet_names}")
        
        all_data = []
        for sheet_name in excel_file.sheet_names:
            logging.info(f"Processing sheet: {sheet_name}")
            df = pd.read_excel(BytesIO(file_bytes), sheet_name=sheet_name)
            logging.info(f"Sheet {sheet_name} shape: {df.shape}")
            logging.info(f"Sheet {sheet_name} columns: {list(df.columns)}")
            all_data.append(df)
        
        if all_data:
            combined_df = pd.concat(all_data, ignore_index=True)
            logging.info(f"Combined data shape: {combined_df.shape}")
            return combined_df
        else:
            logging.warning("No data found in Excel file")
            return pd.DataFrame()
    except Exception as e:
        logging.error(f"Error extracting data from Excel: {e}", exc_info=True)
        return pd.DataFrame()

def validate_and_normalize_tickers(tickers: List[str]) -> Dict[str, str]:
    """
    Use AI to validate and normalize ticker symbols for Alpha Vantage compatibility.
    Returns a mapping of original ticker to normalized ticker.
    """
    if not tickers:
        return {}
    
    prompt = f"""
    You are a financial data expert specializing in stock ticker format normalization for the Alpha Vantage API.

    I have the following stock tickers that need to be checked for Alpha Vantage format compatibility:
    {tickers}

    IMPORTANT: Assume ALL tickers are valid. Only make format changes if you know for certain that Alpha Vantage requires a different format.

    ALPHA VANTAGE FORMAT RULES:
    1. Alpha Vantage does NOT support dots (.) in ticker symbols
    2. For tickers with dots, use hyphens (-) instead (e.g., BRK.B → BRK-B)
    3. Convert all tickers to uppercase
    4. Remove any extra spaces or special characters

    KNOWN FORMAT CORRECTIONS (only these specific cases):
    - BRKB → BRK-B (Berkshire Hathaway Class B - Alpha Vantage uses BRK-B)
    - BRKA → BRK-A (Berkshire Hathaway Class A - Alpha Vantage uses BRK-A)
    - BRK.B → BRK-B (convert dots to hyphens)
    - BRK.A → BRK-A (convert dots to hyphens)

    VALIDATION POLICY:
    - Trust that all tickers are valid
    - Only make corrections for known format issues
    - Do NOT mark any ticker as invalid
    - Leave tickers unchanged unless you know they need format correction

    Return ONLY a JSON object with this exact format:
    {{
        "ticker_mappings": {{
            "BRKB": "BRK-B",
            "BRKA": "BRK-A"
        }},
        "corrections": [
            {{
                "original": "BRKB",
                "corrected": "BRK-B",
                "reason": "Berkshire Hathaway Class B uses BRK-B in Alpha Vantage"
            }}
        ]
    }}

    Only include tickers that need known format corrections. If a ticker doesn't need correction, don't include it in the mappings.
    """
    
    try:
        logging.info(f"Validating and normalizing {len(tickers)} tickers...")
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are a financial data expert specializing in stock ticker validation and normalization for the Alpha Vantage API. Always return valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content
        if content is None:
            logging.error("OpenAI returned None content for ticker validation")
            return {}
        
        result = json.loads(content)
        ticker_mappings = result.get("ticker_mappings", {})
        corrections = result.get("corrections", [])
        
        # Log corrections
        if corrections:
            logging.info(f"Ticker corrections made: {corrections}")
        
        return ticker_mappings
        
    except Exception as e:
        logging.error(f"Error validating tickers: {e}", exc_info=True)
        return {}

def extract_portfolio_with_ai(content: str, file_type: str) -> Dict[str, float]:
    """Use GPT to extract portfolio holdings and validate tickers in a batch."""
    logging.info(f"Starting AI portfolio extraction for {file_type} file...")
    
    # Enhanced prompt for better CSV parsing
    if file_type == 'csv':
        prompt = f"""
        Analyze this CSV portfolio data and extract stock holdings. Look for:
        - Stock ticker symbols (like AAPL, MSFT, GOOGL, etc.)
        - Number of shares, quantities, or positions
        - Common column names: Ticker, Symbol, Stock, Shares, Quantity, Position, Amount
        
        CSV Content:
        {content[:4000]}
        
        Return ONLY a JSON object with this exact format:
        {{
            "holdings": [
                {{"ticker": "AAPL", "shares": 100}},
                {{"ticker": "MSFT", "shares": 50}}
            ]
        }}
        
        Important:
        - Extract ALL stock tickers found in the data
        - Use the number of shares/quantity from the data
        - If no shares found, use 100 as default
        - Convert tickers to uppercase
        - Only include valid stock symbols (3-5 letters)
        """
    else:
        prompt = f"""
        Analyze the following {file_type} content and extract stock portfolio information.
        Extract all stock tickers and the number of shares held. Look for:
        - Stock symbols (like AAPL, MSFT, GOOGL, etc.)
        - Company names that can be mapped to tickers
        - Number of shares, quantities, or positions
        Content:
        {content[:4000]}
        Return the data as a JSON object with this exact format:
        {{
            "holdings": [
                {{"ticker": "AAPL", "shares": 100}},
                {{"ticker": "MSFT", "shares": 50}}
            ]
        }}
        """
    
    try:
        logging.info("Sending request to OpenAI for portfolio analysis...")
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "You are a financial analyst expert at extracting portfolio data from documents. Always return valid JSON with stock tickers and share quantities."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        
        logging.info("Received response from OpenAI, parsing JSON...")
        content = response.choices[0].message.content
        if content is None:
            logging.error("OpenAI returned None content")
            return {}
        result = json.loads(content)  # type: ignore
        logging.info(f"OpenAI extracted result: {result}")
        
        # Extract initial holdings
        initial_holdings = {
            item.get("ticker", "").upper(): float(item.get("shares", 100))
            for item in result.get("holdings", []) if item.get("ticker")
        }
        
        logging.info(f"Initial extracted holdings: {initial_holdings}")
        
        if not initial_holdings:
            logging.warning("No potential holdings found in AI response")
            return {}

        # Validate and normalize tickers
        ticker_list = list(initial_holdings.keys())
        ticker_mappings = validate_and_normalize_tickers(ticker_list)
        
        # Apply corrections and create final holdings
        final_holdings = {}
        corrections_made = []
        
        for original_ticker, shares in initial_holdings.items():
            # Get normalized ticker, default to original if not in mappings
            if original_ticker in ticker_mappings:
                mapped_ticker = ticker_mappings[original_ticker]
                normalized_ticker = mapped_ticker if mapped_ticker and mapped_ticker != "" else original_ticker
            else:
                normalized_ticker = original_ticker
            
            # Apply correction if different from original
            if normalized_ticker != original_ticker:
                corrections_made.append(f"{original_ticker} → {normalized_ticker}")
                logging.info(f"Ticker correction: {original_ticker} → {normalized_ticker}")
                final_holdings[normalized_ticker] = shares
            else:
                # No correction needed, use original ticker
                final_holdings[original_ticker] = shares
        
        if corrections_made:
            logging.info(f"Ticker corrections applied: {corrections_made}")
        
        logging.info(f"Final normalized holdings: {final_holdings}")
        return final_holdings
        
    except Exception as e:
        logging.error(f"Error extracting portfolio with AI: {e}", exc_info=True)
        return {}

def process_single_file(file, file_type: str) -> Dict[str, float]:
    """Process a single file and return extracted holdings."""
    try:
        file_bytes = file.read()
        content = ""
        
        if file_type == 'pdf':
            content = extract_text_from_pdf(file_bytes)
        elif file_type == 'docx':
            content = extract_text_from_docx(file_bytes)
        elif file_type == 'csv':
            df = pd.read_csv(BytesIO(file_bytes))
            content = df.to_string()
        elif file_type in ['xlsx', 'xls']:
            df = extract_data_from_excel(file_bytes)
            content = df.to_string()
        
        if not content:
            logging.warning(f"No content extracted from {file.name}")
            return {}
        
        holdings = extract_portfolio_with_ai(content, file_type)
        return holdings
        
    except Exception as e:
        logging.error(f"Error processing file {file.name}: {e}")
        return {}

def combine_holdings(holdings_list: List[Dict[str, float]]) -> Dict[str, float]:
    """Combine multiple holdings dictionaries, summing shares for duplicate tickers."""
    combined = {}
    for holdings in holdings_list:
        for ticker, shares in holdings.items():
            if ticker in combined:
                combined[ticker] += shares
            else:
                combined[ticker] = shares
    return combined

# ---------- Streamlit UI ----------
def main():
    st.set_page_config(
        page_title="Wall Street Weekly - Portfolio Management",
        page_icon="💼",
        layout="wide"
    )

    if 'google_sheet_initialized' not in st.session_state:
        st.session_state['google_sheet_initialized'] = init_google_sheet()

    if not st.session_state['google_sheet_initialized']:
        st.error("Failed to initialize database connection. Please check system configuration.")
        return

    st.markdown("""
    <style>
    .main-header { 
        text-align: center; 
        padding: 3rem 0; 
        background: linear-gradient(135deg, #1a365d 0%, #2d3748 100%); 
        color: white; 
        border-radius: 12px; 
        margin-bottom: 2.5rem; 
        box-shadow: 0 8px 32px rgba(0,0,0,0.15);
        border: 1px solid rgba(255,255,255,0.1);
    }
    .main-header h1 {
        font-size: 2.8rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        text-shadow: 0 2px 4px rgba(0,0,0,0.3);
        letter-spacing: -0.5px;
    }
    .main-header h3 {
        font-size: 1.3rem;
        font-weight: 400;
        opacity: 0.9;
        margin: 0;
        letter-spacing: 0.5px;
    }
    .upload-section, .portfolio-display {
        background: #fff;
        padding: 2.5rem;
        border-radius: 12px;
        border: 1px solid #e2e8f0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        margin-bottom: 2rem;
    }
    .section-title {
        font-size: 1.5rem;
        font-weight: 600;
        color: #1a365d;
        margin-bottom: 1rem;
        border-bottom: 2px solid #e2e8f0;
        padding-bottom: 0.5rem;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="main-header"><h1>Wall Street Weekly</h1><h3>Professional Portfolio Management System</h3></div>', unsafe_allow_html=True)

    # Contact Advisor button at the top (centered)
    st.markdown("<div style='display: flex; justify-content: center; margin-bottom: 2rem;'><button style='background: linear-gradient(135deg, #1a365d 0%, #2d3748 100%); color: white; border: none; border-radius: 8px; padding: 0.75rem 2rem; font-weight: 600; font-size: 1rem; cursor: pointer;' onclick=\"window.location.href='mailto:keanejpalmer@gmail.com'\">Contact Financial Advisor</button></div>", unsafe_allow_html=True)

    st.markdown("---")

    # Main content: two columns for upload and overview
    col1, col2 = st.columns([1, 1], gap="large")

    with col1:
        st.markdown('<div class="upload-section">', unsafe_allow_html=True)
        st.markdown('<div class="section-title">Portfolio Upload</div>', unsafe_allow_html=True)
        st.markdown("Upload your portfolio documents to receive personalized financial newsletters with market analysis and performance insights.")
        email = st.text_input("Email Address", placeholder="you@example.com")
        uploaded_files = st.file_uploader(
            "Upload Portfolio Documents (PDF, DOCX, XLSX, CSV)",
            type=['pdf', 'docx', 'xlsx', 'xls', 'csv'],
            help="Select one or more files containing your portfolio holdings. Supported formats: PDF, Word, Excel, CSV",
            accept_multiple_files=True,
            key="portfolio_uploader"
        )
        
        if uploaded_files:
            st.success(f"Files uploaded: {len(uploaded_files)} file(s)")
            for i, file in enumerate(uploaded_files):
                st.info(f"File {i+1}: {file.name} ({file.size} bytes)")
        else:
            st.info("Please select one or more files to upload")

        # Always show the button, but disable it if inputs are missing
        button_disabled = not (uploaded_files and email)
        process_clicked = st.button("Process Portfolio Files", type="primary", key="process_portfolio", disabled=button_disabled)
        if button_disabled:
            st.info("Please upload files and enter your email to enable processing.")

        if process_clicked and uploaded_files and email:
            try:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                def update_progress(step, total_steps, message):
                    progress = step / total_steps
                    progress_bar.progress(progress)
                    status_text.text(f"Step {step}/{total_steps}: {message}")
                
                # Calculate total steps: 1 for initialization + 1 per file for processing + 1 for saving
                total_steps = 1 + len(uploaded_files) + 1
                current_step = 0
                
                update_progress(current_step, total_steps, "Initializing file processing...")
                current_step += 1
                
                all_holdings = []
                file_results = []
                
                # Process each file individually
                for i, file in enumerate(uploaded_files):
                    file_type = file.name.split('.')[-1].lower()
                    update_progress(current_step, total_steps, f"Processing file {i+1}/{len(uploaded_files)}: {file.name}")
                    
                    st.info(f"Processing: {file.name} ({file.size} bytes)")
                    
                    # Process the file
                    file_holdings = process_single_file(file, file_type)
                    
                    if file_holdings:
                        st.success(f"✓ {file.name}: Found {len(file_holdings)} holdings")
                        for ticker, shares in file_holdings.items():
                            st.info(f"   • {ticker}: {shares} shares")
                        all_holdings.append(file_holdings)
                        file_results.append({
                            'file': file.name,
                            'holdings': file_holdings,
                            'status': 'success'
                        })
                    else:
                        st.warning(f"⚠ {file.name}: No valid holdings found")
                        file_results.append({
                            'file': file.name,
                            'holdings': {},
                            'status': 'no_holdings'
                        })
                    
                    current_step += 1
                
                # Combine all holdings
                update_progress(current_step, total_steps, "Combining portfolio data...")
                combined_holdings = combine_holdings(all_holdings)
                
                if combined_holdings:
                    st.success(f"Combined portfolio: {len(combined_holdings)} unique holdings")
                    for ticker, shares in combined_holdings.items():
                        st.info(f"   • {ticker}: {shares} shares")
                else:
                    st.error("No valid holdings found in any of the uploaded files.")
                    st.info("Please check that your files contain valid stock portfolio data.")
                    return
                
                current_step += 1
                
                # Save to Google Sheets
                update_progress(current_step, total_steps, "Saving portfolio to database...")
                st.info("Saving combined portfolio data to database...")
                
                if save_user_portfolio_to_sheets(email, combined_holdings):
                    update_progress(current_step, total_steps, "Portfolio processing complete!")
                    st.success("Portfolio saved successfully!")
                    st.session_state['current_holdings'] = combined_holdings
                    st.session_state['current_email'] = email
                    st.session_state['file_results'] = file_results
                    st.success("All files processed and portfolio saved successfully!")
                    st.info("You can now view your combined portfolio on the right side and send a test newsletter.")
                    import time
                    time.sleep(2)
                    st.rerun()
                else:
                    st.error("Failed to save portfolio to database.")
                    st.error("Please check the terminal logs for detailed error information.")
                    
            except Exception as e:
                st.error(f"An error occurred during processing: {str(e)}")
                st.error("Please check the terminal for detailed error logs.")
                logging.error(f"Portfolio processing error: {e}", exc_info=True)
        st.markdown('</div>', unsafe_allow_html=True)

    with col2:
        st.markdown('<div class="section-title">Portfolio Overview</div>', unsafe_allow_html=True)
        if 'current_holdings' in st.session_state:
            holdings = st.session_state['current_holdings']
            email = st.session_state.get('current_email', '')
            file_results = st.session_state.get('file_results', [])
            
            st.markdown('<div class="portfolio-display">', unsafe_allow_html=True)
            st.subheader(f"Combined Portfolio for: {email}")
            
            # Show file processing summary
            if file_results:
                st.markdown("**File Processing Summary:**")
                for result in file_results:
                    if result['status'] == 'success':
                        st.success(f"✓ {result['file']}: {len(result['holdings'])} holdings")
                    else:
                        st.warning(f"⚠ {result['file']}: No holdings found")
                st.markdown("---")
            
            ticker_list = tuple(holdings.keys())
            if ticker_list:
                portfolio_data = []
                for ticker, shares in holdings.items():
                    portfolio_data.append({'Ticker': ticker, 'Shares': shares})
                df = pd.DataFrame(portfolio_data)
                st.dataframe(df, use_container_width=True)
                if st.button("Load Current Prices", use_container_width=True, key="load_prices"):
                    with st.spinner("Fetching current prices..."):
                        portfolio_details = get_batch_stock_data(ticker_list)
                        portfolio_data_with_prices = []
                        total_value = 0
                        for ticker, shares in holdings.items():
                            details = portfolio_details.get(ticker)
                            if details and details.get('current_price') is not None:
                                current_price = details.get('current_price', 0)
                                value = current_price * shares
                                total_value += value
                                portfolio_data_with_prices.append({
                                    'Ticker': ticker,
                                    'Company': details.get('company_name', ticker),
                                    'Shares': shares,
                                    'Current Price': f"${current_price:.2f}",
                                    'Value': f"${value:,.2f}"
                                })
                        if portfolio_data_with_prices:
                            df_with_prices = pd.DataFrame(portfolio_data_with_prices)
                            st.dataframe(df_with_prices, use_container_width=True)
                            st.markdown('<div class="metric-card">', unsafe_allow_html=True)
                            st.metric("Total Portfolio Value", f"${total_value:,.2f}")
                            st.markdown('</div>', unsafe_allow_html=True)
                        else:
                            st.warning("Could not fetch current prices for any holdings")
            if st.button("Send Test Newsletter", use_container_width=True, key="send_test_newsletter"):
                with st.spinner(f"Generating and sending newsletter to {email}..."):
                    if generate_newsletter_for_user(email, holdings):
                        st.success(f"Newsletter sent to {email}!")
                    else:
                        st.error("Failed to send newsletter. Check logs for details.")
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div class="portfolio-display">', unsafe_allow_html=True)
            st.info("Upload portfolio files to see your holdings and send newsletters")
            st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()